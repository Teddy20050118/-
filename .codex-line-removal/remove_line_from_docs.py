from __future__ import annotations

from pathlib import Path
import shutil
import tempfile
import zipfile

from docx import Document
from docx.text.paragraph import Paragraph


ROOT = Path(r"C:\Users\eric7\OneDrive\Desktop\pjt")
DELIVERABLES = ROOT / "deliverables"
WORK = ROOT / ".codex-line-removal"
BACKUP = WORK / "backup"

REQ_PATH = DELIVERABLES / "點餐指南針_AI智慧點餐_系統需求書.docx"
PROP_PATH = DELIVERABLES / "點餐指南針_AI智慧點餐_企劃書.docx"

PRESERVED_PACKAGE_PARTS = (
    "[Content_Types].xml",
    "_rels/.rels",
    "word/_rels/document.xml.rels",
    "word/numbering.xml",
    "word/settings.xml",
    "word/styles.xml",
)


def iter_paragraphs(parent):
    for paragraph in parent.paragraphs:
        yield paragraph
    for table in parent.tables:
        for row in table.rows:
            for cell in row.cells:
                yield from iter_paragraphs(cell)


def replace_in_paragraph(paragraph: Paragraph, old: str, new: str) -> bool:
    changed = False
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            changed = True
    return changed


def replace_exact(paragraph: Paragraph, old: str, new: str) -> bool:
    if paragraph.text.strip() != old:
        return False
    if not paragraph.runs:
        paragraph.add_run(new)
        return True
    paragraph.runs[0].text = new
    for run in paragraph.runs[1:]:
        run._element.getparent().remove(run._element)
    return True


def delete_paragraph(paragraph: Paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def restore_preserved_package_parts(source: Path, output: Path) -> None:
    with zipfile.ZipFile(source, "r") as source_zip:
        preserved = {
            name: source_zip.read(name)
            for name in PRESERVED_PACKAGE_PARTS
        }

    with tempfile.NamedTemporaryFile(
        suffix=".docx", delete=False, dir=output.parent
    ) as temporary:
        temporary_path = Path(temporary.name)

    try:
        with (
            zipfile.ZipFile(output, "r") as current_zip,
            zipfile.ZipFile(temporary_path, "w", zipfile.ZIP_DEFLATED) as new_zip,
        ):
            for item in current_zip.infolist():
                data = preserved.get(item.filename, current_zip.read(item.filename))
                new_zip.writestr(item, data)
        temporary_path.replace(output)
    finally:
        if temporary_path.exists():
            temporary_path.unlink()


def save_preserving_package(doc: Document, source: Path, output: Path) -> None:
    doc.save(output)
    restore_preserved_package_parts(source, output)


def update_requirements() -> None:
    doc = Document(REQ_PATH)

    exact_replacements = {
        "系統範圍：Web 與 LINE 對話入口、Google 菜單擷取、菜單照片辨識與人工確認、標準化菜單資料、偏好與預算蒐集、AI 推薦與規則驗證、評價證據整理、健康檢查與操作紀錄。":
            "系統範圍：響應式 Web 入口、Google 菜單擷取、菜單照片辨識與人工確認、標準化菜單資料、偏好與預算蒐集、AI 推薦與規則驗證、評價證據整理、健康檢查與操作紀錄。",
        "安全性：API 金鑰僅由環境變數載入；LINE webhook 必須以原始請求本文驗證 HMAC-SHA256 簽章；上傳檔案須驗證格式、大小與安全檔名。":
            "安全性：API 金鑰僅由環境變數載入；上傳檔案須驗證格式、大小與安全檔名；錯誤訊息與日誌不得輸出密鑰內容。",
        "可維護性：菜單辨識、推薦、評論、LINE 整合與 API 層應模組化；核心規則須具單元測試，並以一致的 schema v2 儲存已確認菜單。":
            "可維護性：菜單辨識、推薦、評論、Web 介面與 API 層應模組化；核心規則須具單元測試，並以一致的 schema v2 儲存已確認菜單。",
        "系統應提供 Web 與 LINE Bot 兩種使用入口，並在同一後端服務中處理菜單、推薦與對話流程。":
            "系統應提供響應式 Web 使用入口，並在同一後端服務中處理菜單、推薦與對話流程。",
        "系統應提供健康檢查端點，顯示菜單載入、Vision 與 LINE 設定狀態，但不得回傳任何密鑰內容。":
            "系統應提供健康檢查端點，顯示菜單載入與 Vision 設定狀態，但不得回傳任何密鑰內容。",
        "a. 在 Web 或 LINE 選擇餐廳並上傳菜單照片。":
            "a. 在 Web 選擇餐廳並上傳菜單照片。",
    }

    for paragraph in list(iter_paragraphs(doc)):
        for old, new in exact_replacements.items():
            replace_exact(paragraph, old, new)

    functional_table = next(
        table
        for table in doc.tables
        if any("OC-F-001" in cell.text for row in table.rows for cell in row.cells)
    )
    for row in list(functional_table.rows):
        if row.cells[0].text.strip() == "OC-F-013":
            row._tr.getparent().remove(row._tr)
            break

    for paragraph in iter_paragraphs(doc):
        replace_in_paragraph(paragraph, "OC-F-014", "OC-F-013")

    line_use_case = next(
        table
        for table in doc.tables
        if "透過 LINE Bot 完成菜單辨識與推薦" in table.cell(0, 1).text
    )
    line_use_case._element.getparent().remove(line_use_case._element)

    for paragraph in iter_paragraphs(doc):
        replace_in_paragraph(paragraph, "OC-UC004", "OC-UC003")
        replace_in_paragraph(paragraph, "OC-UC005", "OC-UC004")

    output = WORK / "requirements-updated.docx"
    save_preserving_package(doc, REQ_PATH, output)
    output.replace(REQ_PATH)


def update_proposal() -> None:
    doc = Document(PROP_PATH)

    exact_replacements = {
        "從手機菜單照片或公開資料建立可修正的結構化菜單，依預算、偏好與忌口產生點餐建議，再以程式驗證品項、價格與限制條件；支援 Web 與 LINE 現場使用。":
            "從手機菜單照片或公開資料建立可修正的結構化菜單，依預算、偏好與忌口產生點餐建議，再以程式驗證品項、價格與限制條件；支援響應式 Web 現場使用。",
        "點餐指南針服務在「看不懂菜單、選擇太多、限制難表達、資訊可信度不足」的用餐情境。使用者可由網頁或 LINE 上傳紙本菜單照片，系統先辨識店名、分類、品項與價格，再以人工確認機制修正衝突，建立可供運算的正式菜單。":
            "點餐指南針服務在「看不懂菜單、選擇太多、限制難表達、資訊可信度不足」的用餐情境。使用者可由響應式網頁上傳紙本菜單照片，系統先辨識店名、分類、品項與價格，再以人工確認機制修正衝突，建立可供運算的正式菜單。",
        "創新四｜同一能力跨介面：Web 適合視覺化確認，LINE 適合現場拍照與即時對話；兩者共享後端、菜單與推薦規則，減少重複開發。":
            "創新四｜行動優先 Web：單一響應式介面整合拍照上傳、品質預覽、編號修正、確認與推薦流程，降低使用者的操作與安裝門檻。",
        "第一層｜互動入口：響應式 Web 提供餐廳選擇、照片辨識、品質預覽與對話推薦；LINE Bot 提供現場拍照、編號修正、確認／取消與推薦對話。":
            "第一層｜互動入口：響應式 Web 提供餐廳選擇、照片辨識、品質預覽、編號修正、確認／取消與對話推薦。",
        "第四層｜資料與外部服務：已確認菜單與評論快取以 JSON schema 保存；外部整合包含 Google 公開資訊、OpenAI-compatible 模型服務及 LINE Messaging API。":
            "第四層｜資料與外部服務：已確認菜單與評論快取以 JSON schema 保存；外部整合包含 Google 公開資訊與 OpenAI-compatible 模型服務。",
        "故障隔離：爬蟲、Vision、推薦模型與 LINE 皆以獨立設定啟用；任一外部服務失敗時，不覆蓋已確認資料，推薦可切換決定式備援。":
            "故障隔離：爬蟲、Vision 與推薦模型皆以獨立設定啟用；任一外部服務失敗時，不覆蓋已確認資料，推薦可切換決定式備援。",
        "LINE 對話採低學習成本指令：使用者可傳照片，依預覽編號輸入「改 3 菜名 …」或「改 3 價格 …」，最後以「確認／取消」完成；自然語言舊句型仍保留相容性。":
            "Web 修正流程採低學習成本設計：使用者可依預覽編號修正菜名或價格，查看衝突與品質指標，最後以「確認／取消」完成。",
        "(4) LINE Messaging API、HTML／CSS／JavaScript、Git 與單元測試":
            "(4) HTML／CSS／JavaScript、Git 與單元測試",
    }

    delete_exact = {
        "LINE webhook 使用原始本文 HMAC-SHA256 驗證；圖片先快速回覆，再於背景處理並推送結果，避免外部模型延遲造成 reply token 逾時。"
    }

    for paragraph in list(iter_paragraphs(doc)):
        if paragraph.text.strip() in delete_exact:
            delete_paragraph(paragraph)
            continue
        for old, new in exact_replacements.items():
            replace_exact(paragraph, old, new)

    output = WORK / "proposal-updated.docx"
    save_preserving_package(doc, PROP_PATH, output)
    output.replace(PROP_PATH)


def main() -> None:
    BACKUP.mkdir(parents=True, exist_ok=True)
    shutil.copy2(REQ_PATH, BACKUP / REQ_PATH.name)
    shutil.copy2(PROP_PATH, BACKUP / PROP_PATH.name)
    update_requirements()
    update_proposal()
    print(REQ_PATH)
    print(PROP_PATH)


if __name__ == "__main__":
    main()
